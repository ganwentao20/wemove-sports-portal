from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "requirements.md"
OUTPUT = ROOT / "docs" / "deliverables" / "下午班-08组-WEMOVE-SPORTS-项目需求文档-v0.2.docx"

NAVY = "17324D"
TEAL = "0D7C86"
LIGHT_TEAL = "E8F4F4"
LIGHT_BLUE = "EEF3F7"
GRAY = "5D6975"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia="等线", latin="Aptos", size=None, bold=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clean_inline(text: str) -> str:
    return re.sub(r"(`|\*\*)", "", text).strip()


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "请在 Word 中更新目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("263746")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.22

    for style_name, size, color in (
        ("Title", 29, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(14 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(7)

    doc.styles["Heading 1"].paragraph_format.page_break_before = True


def configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("WEMOVE SPORTS  ·  软件开发实践2  ·  第8组")
    set_run_font(run, size=8, bold=True, color=TEAL)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Cm(16.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(12.5)
    table.columns[1].width = Cm(4)
    left = table.cell(0, 0).paragraphs[0]
    run = left.add_run("项目需求文档 · v0.2 · 2026-09-06")
    set_run_font(run, size=8, color=GRAY)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run("第 ")
    set_run_font(run, size=8, color=GRAY)
    add_field(right, " PAGE ")
    run = right.add_run(" 页")
    set_run_font(run, size=8, color=GRAY)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("\n\n")
    badge = p.add_run("COURSE PROJECT  ·  GROUP 08")
    set_run_font(badge, size=10, bold=True, color=TEAL)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(10)
    title.add_run("WEMOVE SPORTS\n官网与业务门户重构")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("项目需求文档  ·  Requirements Specification")
    set_run_font(run, size=14, color=GRAY)

    line = doc.add_paragraph()
    line.paragraph_format.space_after = Pt(24)
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, size=9, color=TEAL)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    labels = ["课程", "组别", "版本", "成员"]
    values = [
        "软件开发实践2",
        "下午班 · 第 8 组",
        "v0.2 · 2026-09-06",
        "甘文韬（组长） · 陈婧琳 · 朱容杰 · 周慧莹 · 倪依玲 · 龙祖怡",
    ]
    for idx, (label, value) in enumerate(zip(labels, values)):
        meta.cell(idx, 0).width = Cm(3)
        meta.cell(idx, 1).width = Cm(13)
        set_cell_shading(meta.cell(idx, 0), LIGHT_TEAL)
        for cell in meta.rows[idx].cells:
            set_cell_margins(cell, 140, 160, 140, 160)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r1 = meta.cell(idx, 0).paragraphs[0].add_run(label)
        set_run_font(r1, size=9, bold=True, color=TEAL)
        r2 = meta.cell(idx, 1).paragraphs[0].add_run(value)
        set_run_font(r2, size=9.5, color=NAVY)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(30)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("文档定位")
    set_run_font(run, size=10, bold=True, color=TEAL)
    note = doc.add_paragraph(
        "本文件以课程可验收范围为主线，并把长期产品愿景明确标为 P1/P2；实现状态只以代码、CI、演示和测试证据判定。"
    )
    note.paragraph_format.left_indent = Cm(0.5)

    doc.add_page_break()
    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(12)
    toc_run = toc.add_run("目录")
    set_run_font(toc_run, size=18, bold=True, color=NAVY)
    p = doc.add_paragraph()
    add_field(p, ' TOC \\o "1-3" \\h \\z \\u ')
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(rows):
        tr_pr = table.rows[row_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_idx == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            tr_pr.append(repeat_header)
        for col_idx in range(columns):
            cell = table.cell(row_idx, col_idx)
            text = clean_inline(row[col_idx]) if col_idx < len(row) else ""
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, NAVY)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1.5)
                paragraph.paragraph_format.space_before = Pt(1.5)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=8.3,
                        bold=row_idx == 0,
                        color=WHITE if row_idx == 0 else "263746",
                    )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def render_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    index = 0
    skipped_title = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if skipped_title and stripped.startswith("版本："):
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = []
            for table_line in table_lines:
                cells = [cell.strip() for cell in table_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph(style=f"Heading {level}")
            if level == 1 and heading.group(2).startswith("1."):
                p.paragraph_format.page_break_before = False
            p.add_run(clean_inline(heading.group(2)))
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if numbered or bullet:
            p = doc.add_paragraph(style="Normal" if numbered else "List Bullet")
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            if numbered:
                number = stripped.split('.', 1)[0]
                p.add_run(f"{number}.  {clean_inline(numbered.group(1))}")
            else:
                p.add_run(clean_inline(bullet.group(1)))
            index += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.add_run(clean_inline(stripped))
        index += 1


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])
    add_header_footer(doc.sections[0])
    add_cover(doc)
    render_markdown(doc, SOURCE.read_text(encoding="utf-8"))

    for section in doc.sections:
        configure_page(section)
    doc.core_properties.title = "WEMOVE SPORTS 官网与业务门户重构——项目需求文档"
    doc.core_properties.subject = "软件开发实践2 · 下午班第8组"
    doc.core_properties.author = "甘文韬、陈婧琳、朱容杰、周慧莹、倪依玲、龙祖怡"
    doc.core_properties.keywords = "WEMOVE SPORTS, 软件开发实践2, 需求文档, 第8组"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
